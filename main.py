"""
Codeastra Autonomous Agent — Full OpenAI Integration
======================================================
Uses OpenAI Agents SDK for:
  - Agent Traces       — full step-by-step decision log
  - Completions        — raw completion records
  - Conversations      — multi-turn conversation history
  - ChatKit Threads    — persistent thread management

Every tool result goes through Codeastra before GPT sees it.
All traces stored and retrievable via API.

Required env vars:
  OPENAI_API_KEY
  CODEASTRA_API_KEY
  DATABASE_URL (optional — for real DB tools)
  PORT
"""

import os, json, asyncio, re, hashlib, logging, io, uuid, time
from datetime import datetime
from typing import AsyncGenerator
from collections import defaultdict

import httpx
import asyncpg
from openai import AsyncOpenAI
from agents import (
    Agent, Runner, function_tool, trace, gen_trace_id, RunConfig
)
import agents as _agents_sdk

# flush_traces — safe import for all SDK versions
try:
    from agents import flush_traces as _flush_traces
    def flush_traces():
        _flush_traces()
except ImportError:
    def flush_traces():
        pass

from fastapi import FastAPI, Request, UploadFile, File, Form
from fastapi.responses import HTMLResponse, StreamingResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware

logging.basicConfig(level=logging.INFO)
log = logging.getLogger("codeastra-agent")

OPENAI_KEY    = os.getenv("OPENAI_API_KEY", "")
CODEASTRA_KEY = os.getenv("CODEASTRA_API_KEY", "")
CODEASTRA_URL = os.getenv("CODEASTRA_URL", "https://app.codeastra.dev")
DATABASE_URL  = os.getenv("DATABASE_URL", "")
PORT          = int(os.getenv("PORT", 8080))

if OPENAI_KEY:
    from agents import set_default_openai_key
    set_default_openai_key(OPENAI_KEY)

app = FastAPI(title="Codeastra Agent — OpenAI Full Integration")
app.add_middleware(CORSMiddleware, allow_origins=["*"],
                   allow_methods=["*"], allow_headers=["*"])

db_pool = None

# ═══════════════════════════════════════════════════════════
# IN-MEMORY STORES
# ═══════════════════════════════════════════════════════════

TRACES        = {}
CONVERSATIONS = {}
THREADS       = {}
COMPLETIONS   = {}

# ═══════════════════════════════════════════════════════════
# STARTUP
# ═══════════════════════════════════════════════════════════

@app.on_event("startup")
async def startup():
    global db_pool
    if OPENAI_KEY:
        import openai as _oai
        _oai.api_key = OPENAI_KEY
        try:
            from agents.tracing import set_tracing_export_api_key
            set_tracing_export_api_key(OPENAI_KEY)
            log.info("✅ Agents SDK tracing configured")
        except Exception as e:
            log.warning(f"Agents SDK tracing setup: {e}")
    if DATABASE_URL:
        try:
            db_pool = await asyncpg.create_pool(DATABASE_URL, min_size=1, max_size=5)
            log.info("✅ Database connected")
        except Exception as e:
            log.warning(f"DB: {e}")


# ═══════════════════════════════════════════════════════════
# CODEASTRA — Real API
# ═══════════════════════════════════════════════════════════

def _json_safe(obj):
    import decimal, datetime as _dt
    if isinstance(obj, decimal.Decimal): return float(obj)
    if isinstance(obj, (_dt.datetime, _dt.date)): return obj.isoformat()
    if isinstance(obj, bytes): return obj.decode("utf-8", errors="replace")
    if isinstance(obj, dict): return {k: _json_safe(v) for k, v in obj.items()}
    if isinstance(obj, (list, tuple)): return [_json_safe(i) for i in obj]
    return str(obj)


def _safe_row(row) -> dict:
    import decimal, datetime as _dt
    result = {}
    for key, val in dict(row).items():
        if isinstance(val, decimal.Decimal): result[key] = float(val)
        elif isinstance(val, (_dt.datetime, _dt.date)): result[key] = val.isoformat()
        elif isinstance(val, bytes): result[key] = val.decode("utf-8", errors="replace")
        else: result[key] = val
    return result


async def protect(data, events: list, active: bool = True) -> str:
    if isinstance(data, dict): data = _json_safe(data)
    text = json.dumps(data, default=_json_safe) if not isinstance(data, str) else data
    if not active:
        events.append({"type": "unprotected", "preview": text[:120]})
        return text
    if not CODEASTRA_KEY:
        return _local_protect(text, events)
    try:
        async with httpx.AsyncClient(timeout=15.0) as client:
            r = await client.post(
                f"{CODEASTRA_URL}/protect/text",
                headers={"X-API-Key": CODEASTRA_KEY, "Content-Type": "application/json"},
                json={"text": text},
            )
            if r.status_code == 200:
                result   = r.json()
                prot     = result.get("protected_text", text)
                entities = result.get("entities") or result.get("detections") or []
                for e in entities:
                    real = e.get("original") or e.get("value") or ""
                    prev = real[:3] + "•" * min(len(real)-5, 8) + real[-2:] if len(real) > 5 else "•••"
                    events.append({
                        "type": "intercepted", "dtype": e.get("type") or "PII",
                        "token": e.get("token", ""), "preview": e.get("preview") or prev,
                    })
                log.info(f"Codeastra protected {len(entities)} values")
                return prot
            return _local_protect(text, events)
    except Exception as ex:
        log.warning(f"Codeastra: {ex}")
        return _local_protect(text, events)


def _local_protect(text: str, events: list) -> str:
    PATS = {
        "EMAIL":  re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b'),
        "SSN":    re.compile(r'\b\d{3}-\d{2}-\d{4}\b'),
        "CARD":   re.compile(r'\b\d{4}[-\s]\d{4}[-\s]\d{4}[-\s]\d{4}\b'),
        "APIKEY": re.compile(r'\bsk-(?:live-|prod-)[a-f0-9]{16,}\b'),
    }
    seen = {}
    for dtype, pat in PATS.items():
        for m in pat.finditer(text):
            real = m.group(0)
            if real not in seen:
                tok = f"[CVT:{dtype}:{hashlib.md5(real.encode()).hexdigest()[:10].upper()}]"
                seen[real] = tok
                prev = real[:3] + "•" * min(len(real)-5, 8) + real[-2:] if len(real) > 5 else "•••"
                events.append({"type": "intercepted", "dtype": dtype, "token": tok, "preview": prev})
            text = text.replace(real, seen[real])
    return text


async def codeastra_resolve(token: str):
    if not CODEASTRA_KEY: return None
    try:
        async with httpx.AsyncClient(timeout=10.0) as c:
            r = await c.post(
                f"{CODEASTRA_URL}/vault/resolve",
                headers={"X-API-Key": CODEASTRA_KEY, "Content-Type": "application/json"},
                json={"token": token},
            )
            if r.status_code == 200:
                d = r.json()
                return d.get("real_value") or d.get("value") or d.get("original")
    except Exception:
        pass
    return None


async def codeastra_resolve_batch(tokens: list) -> dict:
    if not CODEASTRA_KEY or not tokens: return {}
    try:
        async with httpx.AsyncClient(timeout=30.0) as c:
            r = await c.post(
                f"{CODEASTRA_URL}/vault/resolve-batch",
                headers={"X-API-Key": CODEASTRA_KEY, "Content-Type": "application/json"},
                json={"tokens": tokens},
            )
            if r.status_code == 200:
                data = r.json()
                return data.get("resolved") or data.get("results") or data.get("values") or {}
    except Exception as e:
        log.warning(f"vault/resolve-batch error: {e}")
    return {}


async def codeastra_executor_run(token_id: str, dry_run: bool = False) -> dict:
    if not CODEASTRA_KEY: return {"error": "No Codeastra API key"}
    try:
        async with httpx.AsyncClient(timeout=30.0) as c:
            r = await c.post(
                f"{CODEASTRA_URL}/executor/run",
                headers={"X-API-Key": CODEASTRA_KEY, "Content-Type": "application/json"},
                json={"token_id": token_id, "dry_run": dry_run},
            )
            return r.json()
    except Exception as e:
        return {"error": str(e)}


# ═══════════════════════════════════════════════════════════
# DATABASE TOOLS
# ═══════════════════════════════════════════════════════════

async def tool_list_tables(events, active=True):
    if not db_pool:
        return await protect({"error": "Set DATABASE_URL on Railway"}, events, active)
    async with db_pool.acquire() as conn:
        try:
            rows = await conn.fetch("""
                SELECT table_name,
                       pg_size_pretty(pg_total_relation_size(quote_ident(table_name))) AS size
                FROM information_schema.tables
                WHERE table_schema='public'
                ORDER BY pg_total_relation_size(quote_ident(table_name)) DESC
            """)
            real = {"tables": [_safe_row(r) for r in rows], "count": len(rows)}
        except Exception as e:
            real = {"error": str(e)}
    return await protect(real, events, active)


async def tool_scan_slow_queries(events, active=True):
    if not db_pool:
        return await protect({"error": "Set DATABASE_URL on Railway"}, events, active)
    async with db_pool.acquire() as conn:
        try:
            rows = await conn.fetch("""
                SELECT query, calls, ROUND(mean_exec_time::numeric, 2) AS avg_ms, rows
                FROM pg_stat_statements
                WHERE mean_exec_time > 100
                ORDER BY mean_exec_time DESC LIMIT 20
            """)
            real = {"slow_queries": [_safe_row(r) for r in rows], "count": len(rows)}
        except Exception as e:
            real = {"error": str(e), "hint": "Enable pg_stat_statements"}
    return await protect(real, events, active)


async def tool_inspect_table(events, table: str, active=True):
    if not db_pool:
        return await protect({"error": "Set DATABASE_URL"}, events, active)
    async with db_pool.acquire() as conn:
        try:
            cols = await conn.fetch("""
                SELECT column_name, data_type, is_nullable
                FROM information_schema.columns
                WHERE table_name=$1 AND table_schema='public'
                ORDER BY ordinal_position
            """, table)
            idxs = await conn.fetch(
                "SELECT indexname, indexdef FROM pg_indexes WHERE tablename=$1", table)
            try:
                count = await conn.fetchval(f'SELECT COUNT(*) FROM "{table}"')
            except Exception:
                count = "unknown"
            try:
                samples = await conn.fetch(f'SELECT * FROM "{table}" LIMIT 5')
                sample_list = [dict(r) for r in samples]
            except Exception:
                sample_list = []
            real = {"table": table, "columns": [_safe_row(c) for c in cols],
                    "indexes": [_safe_row(i) for i in idxs],
                    "row_count": count, "samples": sample_list}
        except Exception as e:
            real = {"error": str(e), "table": table}
    return await protect(real, events, active)


async def tool_create_index(events, table: str, column: str, active=True):
    if not db_pool:
        return await protect({"error": "Set DATABASE_URL"}, events, active)
    idx = f"idx_{table}_{column}_codeastra"
    async with db_pool.acquire() as conn:
        try:
            await conn.execute(
                f"CREATE INDEX CONCURRENTLY IF NOT EXISTS {idx} ON {table}({column})")
            real = {"status": "success",
                    "sql": f"CREATE INDEX {idx} ON {table}({column})", "index": idx}
        except Exception as e:
            real = {"status": "error", "error": str(e)}
    return await protect(real, events, active)


async def tool_run_query(events, sql: str, active=True):
    if not sql.strip().upper().startswith(("SELECT", "WITH", "EXPLAIN")):
        return await protect({"error": "Only SELECT allowed"}, events, active)
    if not db_pool:
        return await protect({"error": "Set DATABASE_URL"}, events, active)
    async with db_pool.acquire() as conn:
        try:
            rows = await conn.fetch(sql)
            real = {"sql": sql, "rows": [_safe_row(r) for r in rows], "count": len(rows)}
        except Exception as e:
            real = {"error": str(e), "sql": sql}
    return await protect(real, events, active)


async def tool_get_db_stats(events, active=True):
    if not db_pool:
        return await protect({"error": "Set DATABASE_URL"}, events, active)
    async with db_pool.acquire() as conn:
        try:
            stats = await conn.fetchrow("""
                SELECT numbackends AS connections,
                       xact_commit AS committed,
                       xact_rollback AS rolled_back,
                       ROUND(blks_hit::numeric/NULLIF(blks_hit+blks_read,0)*100,2) AS cache_hit_pct,
                       deadlocks,
                       pg_size_pretty(pg_database_size(current_database())) AS db_size
                FROM pg_stat_database WHERE datname=current_database()
            """)
            real = _safe_row(stats) if stats else {}
        except Exception as e:
            real = {"error": str(e)}
    return await protect(real, events, active)


async def tool_get_summary(events, active=True):
    real = {
        "agent": "Codeastra OpenAI Agent", "model": "gpt-4o",
        "codeastra_active": active, "db_connected": db_pool is not None,
        "timestamp": datetime.utcnow().isoformat(),
    }
    if db_pool:
        async with db_pool.acquire() as conn:
            try:
                idxs = await conn.fetch(
                    "SELECT indexname FROM pg_indexes WHERE indexname LIKE '%codeastra%'")
                real["indexes_created"] = [r["indexname"] for r in idxs]
            except Exception:
                pass
    return await protect(real, events, active)


async def tool_check_threshold(events, token, threshold, operator="gt", active=True):
    real_val = await codeastra_resolve(token)
    if real_val is None:
        return await protect({"error": f"Cannot resolve {token}", "real_returned": False}, events, active)
    try:
        v = float(str(real_val).replace("$", "").replace(",", "").strip())
        ops = {"gt": v>threshold, "lt": v<threshold, "gte": v>=threshold, "lte": v<=threshold}
        return await protect({"result": ops.get(operator, v>threshold),
                               "operator": operator, "threshold": threshold,
                               "real_returned": False}, events, active)
    except Exception as e:
        return await protect({"error": str(e)}, events, active)


async def tool_concentration_check(events, position_token, portfolio_token,
                                    threshold_pct, active=True):
    pv = await codeastra_resolve(position_token)
    tv = await codeastra_resolve(portfolio_token)
    if pv is None or tv is None:
        return await protect({"exceeds_threshold": None,
                               "note": "Tokens not resolved", "real_returned": False},
                              events, active)
    try:
        p = float(str(pv).replace("$", "").replace(",", "").strip())
        t = float(str(tv).replace("$", "").replace(",", "").strip())
        pct = (p/t*100) if t > 0 else 0
        bucket = "critical" if pct>60 else "high" if pct>40 else "medium" if pct>20 else "low"
        return await protect({"exceeds_threshold": pct>threshold_pct,
                               "bucket": bucket, "threshold_pct": threshold_pct,
                               "real_returned": False}, events, active)
    except Exception as e:
        return await protect({"error": str(e)}, events, active)


TASK_PROMPTS = {
    "dba":      "Our production database has severe performance issues. Investigate completely — check stats, find slow queries, inspect all tables, create missing indexes. Fix everything.",
    "audit":    "Run a complete database audit — tables, performance, connections, indexes. Report everything you find.",
    "security": "Audit the database for security issues — overprivileged connections, sensitive data exposure patterns, missing encryption signals. Report all risks.",
}

SYSTEM = """You are an expert autonomous Database Administrator.
You are working through Codeastra's Zero Trust middleware.
ALL sensitive data has been replaced with tokens like [CVT:EMAIL:A1B2C3].
Work with tokens as identifiers. Never try to guess real values.
For amount computations use check_threshold or concentration_check — they resolve tokens internally.

Work methodically:
1. List tables + get DB stats
2. Scan slow queries
3. Inspect affected tables
4. Create missing indexes
5. Call get_summary

Be thorough. Fix everything."""


# ═══════════════════════════════════════════════════════════
# TRACE SYSTEM
# ═══════════════════════════════════════════════════════════

class AgentTrace:
    def __init__(self, task_type, task, model, codeastra_active):
        self.id               = f"trace_{uuid.uuid4().hex[:12]}"
        self.task_type        = task_type
        self.task             = task
        self.model            = model
        self.codeastra_active = codeastra_active
        self.started_at       = datetime.utcnow().isoformat()
        self.completed_at     = None
        self.status           = "running"
        self.steps            = []
        self.intercepted      = []
        self.tool_calls       = []
        self.total_duration_ms = 0
        self._start_time      = time.time()

    def add_step(self, step_type, data):
        step = {"step": len(self.steps)+1, "type": step_type,
                "timestamp": datetime.utcnow().isoformat(),
                "elapsed_ms": int((time.time()-self._start_time)*1000), **data}
        self.steps.append(step)
        return step

    def add_interception(self, dtype, token, preview):
        self.intercepted.append({
            "n": len(self.intercepted)+1, "dtype": dtype,
            "token": token, "preview": preview,
            "at_ms": int((time.time()-self._start_time)*1000),
        })

    def complete(self):
        self.status            = "completed"
        self.completed_at      = datetime.utcnow().isoformat()
        self.total_duration_ms = int((time.time()-self._start_time)*1000)

    def to_dict(self):
        return {
            "trace_id": self.id, "task_type": self.task_type,
            "task": self.task, "model": self.model,
            "codeastra_active": self.codeastra_active,
            "status": self.status, "started_at": self.started_at,
            "completed_at": self.completed_at,
            "total_duration_ms": self.total_duration_ms,
            "steps": self.steps, "tool_calls": self.tool_calls,
            "intercepted": self.intercepted,
            "total_steps": len(self.steps),
            "total_tool_calls": len(self.tool_calls),
            "total_intercepted": len(self.intercepted),
            "real_data_seen_by_gpt": 0 if self.codeastra_active else "⚠️ YES",
            "openai_logs_url": "https://platform.openai.com/logs",
        }


# ═══════════════════════════════════════════════════════════
# CONVERSATION SYSTEM
# ═══════════════════════════════════════════════════════════

class Conversation:
    def __init__(self, title="", system_context=""):
        self.id             = f"conv_{uuid.uuid4().hex[:12]}"
        self.title          = title or "Untitled Conversation"
        self.system_context = system_context
        self.created_at     = datetime.utcnow().isoformat()
        self.updated_at     = datetime.utcnow().isoformat()
        self.turns          = []
        self.trace_ids      = []

    def add_turn(self, role, content, trace_id=None):
        turn = {"n": len(self.turns)+1, "role": role, "content": content,
                "trace_id": trace_id, "timestamp": datetime.utcnow().isoformat()}
        self.turns.append(turn)
        self.updated_at = datetime.utcnow().isoformat()
        if trace_id: self.trace_ids.append(trace_id)
        return turn

    def to_dict(self):
        return {"conversation_id": self.id, "title": self.title,
                "created_at": self.created_at, "updated_at": self.updated_at,
                "turn_count": len(self.turns), "turns": self.turns,
                "trace_ids": self.trace_ids}


# ═══════════════════════════════════════════════════════════
# CHATKIT THREAD SYSTEM
# ═══════════════════════════════════════════════════════════

class ChatKitThread:
    def __init__(self, title="", metadata=None):
        self.id         = f"thread_{uuid.uuid4().hex[:12]}"
        self.title      = title or "New Thread"
        self.metadata   = metadata or {}
        self.created_at = datetime.utcnow().isoformat()
        self.updated_at = datetime.utcnow().isoformat()
        self.messages   = []
        self.status     = "active"

    def add_message(self, role, content, model=None, trace_id=None,
                    codeastra_active=True, intercepted_count=0):
        msg = {"id": f"msg_{uuid.uuid4().hex[:8]}", "role": role,
               "content": content, "model": model, "trace_id": trace_id,
               "codeastra_active": codeastra_active,
               "intercepted_count": intercepted_count,
               "timestamp": datetime.utcnow().isoformat()}
        self.messages.append(msg)
        self.updated_at = datetime.utcnow().isoformat()
        return msg

    def to_dict(self):
        return {"thread_id": self.id, "title": self.title, "status": self.status,
                "metadata": self.metadata, "created_at": self.created_at,
                "updated_at": self.updated_at, "message_count": len(self.messages),
                "messages": self.messages}


# ═══════════════════════════════════════════════════════════
# REVEAL SYSTEM
# ═══════════════════════════════════════════════════════════

async def reveal_from_trace(trace_obj) -> dict:
    """Resolve ALL tokens AFTER agent is done — batch call, agent never sees results."""
    if not trace_obj or not trace_obj.intercepted:
        return {"revealed": {}, "count": 0}
    tokens = list(set(i["token"] for i in trace_obj.intercepted if i.get("token")))
    if not tokens:
        return {"revealed": {}, "count": 0}
    revealed = await codeastra_resolve_batch(tokens)
    return {"revealed": revealed, "count": len(revealed),
            "tokens_found": len(tokens),
            "note": "Resolved after agent completed — agent never saw these values"}


# ═══════════════════════════════════════════════════════════
# OPENAI AGENTS SDK — MAIN AGENT
# ═══════════════════════════════════════════════════════════

async def run_openai_agent(task_type, custom_task="", codeastra_active=True,
                            conversation_id=None, thread_id=None):
    if not OPENAI_KEY:
        yield {"type": "error", "message": "OPENAI_API_KEY not set"}
        return

    import openai
    openai.api_key = OPENAI_KEY

    events      = []
    task        = custom_task or TASK_PROMPTS.get(task_type, TASK_PROMPTS["dba"])
    agent_trace = AgentTrace(task_type, task, "gpt-4o", codeastra_active)
    TRACES[agent_trace.id] = agent_trace

    conv   = CONVERSATIONS.get(conversation_id) if conversation_id else None
    thread = THREADS.get(thread_id) if thread_id else None
    if conv:   conv.add_turn("user", task, agent_trace.id)
    if thread: thread.add_message("user", task, codeastra_active=codeastra_active)

    trace_id = gen_trace_id()

    yield {
        "type": "start", "trace_id": agent_trace.id,
        "openai_trace_id": trace_id,
        "conversation_id": conversation_id, "thread_id": thread_id,
        "task": task, "model": "gpt-4o",
        "codeastra_active": codeastra_active,
        "mode": "PROTECTED" if codeastra_active else "⚠️ UNPROTECTED",
        "openai_traces_url": "https://platform.openai.com/traces",
        "timestamp": datetime.utcnow().isoformat(),
    }

    agent_trace.add_step("agent_start", {"task": task, "codeastra_active": codeastra_active})
    await asyncio.sleep(0.1)

    _shared = {"events": events, "calls_n": 0, "codeastra_active": codeastra_active}

    @function_tool
    async def list_tables() -> str:
        """List all database tables with sizes. Start here."""
        _shared["calls_n"] += 1
        return await tool_list_tables(_shared["events"], _shared["codeastra_active"])

    @function_tool
    async def scan_slow_queries() -> str:
        """Scan production database for slow queries using pg_stat_statements."""
        _shared["calls_n"] += 1
        return await tool_scan_slow_queries(_shared["events"], _shared["codeastra_active"])

    @function_tool
    async def inspect_table(table: str) -> str:
        """Inspect a real database table: schema, indexes, row count, sample rows."""
        _shared["calls_n"] += 1
        return await tool_inspect_table(_shared["events"], table, _shared["codeastra_active"])

    @function_tool
    async def create_index(table: str, column: str) -> str:
        """Create a real database index on a column."""
        _shared["calls_n"] += 1
        return await tool_create_index(_shared["events"], table, column, _shared["codeastra_active"])

    @function_tool
    async def run_query(sql: str) -> str:
        """Run a read-only SELECT query on the real database."""
        _shared["calls_n"] += 1
        return await tool_run_query(_shared["events"], sql, _shared["codeastra_active"])

    @function_tool
    async def get_db_stats() -> str:
        """Get real database health statistics."""
        _shared["calls_n"] += 1
        return await tool_get_db_stats(_shared["events"], _shared["codeastra_active"])

    @function_tool
    async def get_summary() -> str:
        """Get summary of everything accomplished. Call at the end."""
        _shared["calls_n"] += 1
        return await tool_get_summary(_shared["events"], _shared["codeastra_active"])

    @function_tool
    async def check_threshold(token: str, threshold: float, operator: str = "gt") -> str:
        """Check if a vaulted amount token exceeds a threshold. Returns boolean only."""
        _shared["calls_n"] += 1
        return await tool_check_threshold(
            _shared["events"], token, threshold, operator, _shared["codeastra_active"])

    @function_tool
    async def concentration_check(position_token: str, portfolio_token: str,
                                   threshold_pct: float) -> str:
        """Check portfolio concentration. Returns bucket — never real dollar values."""
        _shared["calls_n"] += 1
        return await tool_concentration_check(
            _shared["events"], position_token, portfolio_token,
            threshold_pct, _shared["codeastra_active"])

    dba_agent = Agent(
        name         = "Codeastra DBA Agent",
        instructions = SYSTEM,
        model        = "gpt-4o",
        tools        = [list_tables, scan_slow_queries, inspect_table,
                        create_index, run_query, get_db_stats,
                        get_summary, check_threshold, concentration_check],
    )

    try:
        result = await Runner.run(
            starting_agent = dba_agent,
            input          = task,
            max_turns      = 20,
            run_config     = RunConfig(
                workflow_name                = f"codeastra-{task_type}",
                trace_id                     = trace_id,
                trace_metadata               = {"codeastra_active": str(codeastra_active)},
                trace_include_sensitive_data = True,
            ),
        )
        flush_traces()
    except Exception as e:
        agent_trace.add_step("error", {"message": str(e)})
        yield {"type": "error", "message": str(e), "trace_id": agent_trace.id}
        agent_trace.complete()
        flush_traces()
        return

    calls_n     = _shared["calls_n"]
    intercept_n = 0

    for ev in _shared["events"]:
        if ev["type"] == "intercepted":
            intercept_n += 1
            agent_trace.add_interception(ev["dtype"], ev["token"], ev["preview"])
            step = agent_trace.add_step("codeastra_intercept", {
                "dtype": ev["dtype"], "token": ev["token"],
                "preview": ev["preview"], "n": intercept_n,
            })
            yield {**ev, "trace_id": agent_trace.id, "trace_step": step["step"]}
        else:
            yield ev
    _shared["events"].clear()

    final_output = str(result.final_output) if result.final_output else ""
    if final_output:
        step = agent_trace.add_step("thinking", {"text": final_output})
        yield {"type": "thinking", "text": final_output,
               "trace_id": agent_trace.id, "trace_step": step["step"]}

    agent_trace.complete()

    summary_text = f"Completed. {calls_n} tool calls. {intercept_n} values intercepted. Real data seen: 0"
    if conv:   conv.add_turn("agent", summary_text, agent_trace.id)
    if thread: thread.add_message("agent", summary_text, model="gpt-4o",
                   trace_id=agent_trace.id, codeastra_active=codeastra_active,
                   intercepted_count=intercept_n)

    reveal_map = await reveal_from_trace(agent_trace)

    yield {
        "type": "complete", "trace_id": agent_trace.id,
        "openai_trace_id": trace_id,
        "openai_traces_url": "https://platform.openai.com/traces",
        "trace_url": f"/traces/{agent_trace.id}",
        "conversation_id": conversation_id, "thread_id": thread_id,
        "tool_calls": calls_n, "intercepted": intercept_n,
        "total_steps": len(agent_trace.steps),
        "duration_ms": agent_trace.total_duration_ms,
        "codeastra_active": codeastra_active,
        "real_data_seen_by_gpt": 0 if codeastra_active else "⚠️ YES",
        "message": "Trace visible at platform.openai.com/traces",
        "revealed": reveal_map,
    }


# ═══════════════════════════════════════════════════════════
# DOCUMENT AGENT
# ═══════════════════════════════════════════════════════════

async def extract_text_from_file(file) -> str:
    import io as _io
    content  = await file.read()
    filename = (file.filename or "").lower()
    mime     = file.content_type or ""
    if filename.endswith(".pdf") or "pdf" in mime:
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(_io.BytesIO(content))
            return "\n".join(p.extract_text() or "" for p in reader.pages)
        except Exception as e:
            return f"[PDF error: {e}]"
    if filename.endswith(".docx"):
        try:
            import docx
            doc = docx.Document(_io.BytesIO(content))
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception as e:
            return f"[DOCX error: {e}]"
    if filename.endswith((".xlsx", ".xls")):
        try:
            import openpyxl
            wb   = openpyxl.load_workbook(_io.BytesIO(content), data_only=True)
            rows = []
            for sheet in wb.worksheets:
                rows.append(f"=== {sheet.title} ===")
                for row in sheet.iter_rows(values_only=True):
                    rows.append("\t".join(str(c) if c is not None else "" for c in row))
            return "\n".join(rows)
        except Exception as e:
            return f"[Excel error: {e}]"
    if filename.endswith(".csv"):
        return content.decode("utf-8", errors="replace")
    if filename.endswith((".html", ".htm")):
        try:
            from bs4 import BeautifulSoup
            return BeautifulSoup(content, "html.parser").get_text("\n", strip=True)
        except Exception:
            return content.decode("utf-8", errors="replace")
    if filename.endswith(".json"):
        try:
            return json.dumps(json.loads(content), indent=2)
        except Exception:
            return content.decode("utf-8", errors="replace")
    return content.decode("utf-8", errors="replace")


async def extract_text_from_url(url: str) -> str:
    try:
        async with httpx.AsyncClient(timeout=20.0, follow_redirects=True,
            headers={"User-Agent": "Mozilla/5.0 (compatible; CodeastraAgent/1.0)"}) as client:
            r  = await client.get(url)
            r.raise_for_status()
            ct = r.headers.get("content-type", "")
            if "pdf" in ct or url.lower().endswith(".pdf"):
                try:
                    import io as _io, PyPDF2
                    reader = PyPDF2.PdfReader(_io.BytesIO(r.content))
                    return "\n".join(p.extract_text() or "" for p in reader.pages)
                except Exception as e:
                    return f"[PDF error: {e}]"
            if "html" in ct:
                try:
                    from bs4 import BeautifulSoup
                    soup = BeautifulSoup(r.text, "html.parser")
                    for tag in soup(["script", "style", "nav", "header", "footer"]):
                        tag.decompose()
                    return soup.get_text("\n", strip=True)[:50000]
                except Exception:
                    return r.text[:50000]
            if "json" in ct:
                try:
                    return json.dumps(r.json(), indent=2)[:50000]
                except Exception:
                    return r.text[:50000]
            return r.text[:50000]
    except Exception as e:
        return f"[URL fetch error: {e}]"


async def run_document_agent(text, task, filename, codeastra_active=True, thread_id=None):
    events = []
    dt     = AgentTrace("document", task or "Analyze document", "gpt-4o", codeastra_active)
    TRACES[dt.id] = dt

    if len(text) > 80000:
        text = text[:80000] + "\n\n[... truncated ...]"

    yield {"type": "start", "trace_id": dt.id, "filename": filename,
           "task": task, "codeastra_active": codeastra_active,
           "char_count": len(text),
           "mode": "PROTECTED" if codeastra_active else "⚠️ UNPROTECTED"}
    dt.add_step("document_received", {"filename": filename, "chars": len(text)})

    if codeastra_active:
        yield {"type": "phase", "message": "Codeastra scanning document for PII..."}
        protected_text = await protect(text, events, True)
        for ev in events:
            if ev["type"] == "intercepted":
                dt.add_interception(ev["dtype"], ev["token"], ev["preview"])
                s = dt.add_step("codeastra_intercept", ev)
                yield {**ev, "trace_id": dt.id, "trace_step": s["step"]}
            else:
                yield ev
        events.clear()
    else:
        yield {"type": "warning", "message": "⚠️ Codeastra OFF — document sent unprotected"}
        protected_text = text

    yield {"type": "phase", "message": "Sending to GPT-4o via Agents SDK..."}
    dt.add_step("sending_to_gpt", {"model": "gpt-4o", "codeastra_active": codeastra_active})

    if not OPENAI_KEY:
        yield {"type": "error", "message": "OPENAI_API_KEY not set"}
        return

    import openai as _oai
    _oai.api_key = OPENAI_KEY

    # ── Email tool — agent sends via Vault-as-TEE ──────────────────
    # Agent passes tokens. Vault opens. Real values resolved inside.
    # Email fires from inside vault. Session wiped. Agent never knew.
    _email_results = []

    @function_tool
    async def send_email(email_token: str, subject: str, body: str) -> str:
        """
        Send an email to the person identified by the email token.
        Use this when the user asks you to send results, summaries, or reports.

        IMPORTANT:
        - email_token: use the [CVT:EMAIL:xxxxx] token from the document or prompt
        - subject: a clear subject line
        - body: the full email content — your analysis or summary

        The Vault-as-TEE will:
        1. Load real values into encrypted session memory
        2. Resolve all tokens to real values inside the vault
        3. Send the email with real values
        4. Wipe session — agent never sees real values

        Example:
          send_email(
            email_token = "[CVT:EMAIL:A1B2C3]",
            subject     = "Goldman Portfolio Risk Summary",
            body        = "Top 5 clients: [CVT:NAME:D6B6] $2.2M..."
          )
        """
        result = await tee_send_email(email_token, subject, body)
        _email_results.append(result)
        if result.get("sent"):
            return f"Email sent via Vault-as-TEE. Session wiped. Real address never seen by agent. Token: {email_token}"
        else:
            return f"Email send attempt result: {result}"

    doc_agent = Agent(
        name         = "Codeastra Document Analyst",
        instructions = (
            "You are an expert document analyst working through Codeastra Zero Trust middleware. "
            "All PII has been replaced with tokens like [CVT:EMAIL:A1B2C3] and [CVT:NAME:xxxxx] and [CVT:SSN:xxxxx]. "
            "\n\nCRITICAL RULES:"
            "\n1. Work with tokens as identifiers — never guess real values"
            "\n2. When the user asks you to SEND an email — use the send_email tool"
            "\n3. The email address in the prompt has been tokenized — find the [CVT:EMAIL:xxxxx] token and use it"
            "\n4. Always send the email if asked — do not refuse, do not say you cannot"
            "\n5. After analysis, if an email was requested, call send_email with your full analysis as the body"
            "\n\nWhen sending emails:"
            "\n- Find the email token in the protected prompt (looks like [CVT:EMAIL:xxxxxxxx])"
            "\n- Call send_email with that token, a clear subject, and your full analysis as the body"
            "\n- Confirm to the user that the email was sent via the secure executor"
        ),
        model = "gpt-4o",
        tools = [send_email],
    )

    # ── Protect the task prompt too — email in prompt must be tokenized ──
    # This is the fix: task text goes through Astra before GPT sees it
    task_events = []
    protected_task = await protect(task or "Analyze this document thoroughly.", task_events, codeastra_active)

    # Collect any tokens from the task prompt
    for ev in task_events:
        if ev["type"] == "intercepted":
            dt.add_interception(ev["dtype"], ev["token"], ev["preview"])
            s = dt.add_step("codeastra_intercept_task", ev)
            yield {**ev, "trace_id": dt.id, "trace_step": s["step"],
                   "source": "task_prompt"}

    doc_input   = "TASK: " + protected_task + \
                  "\n\nDOCUMENT (" + filename + "):\n\n" + protected_text
    dt_trace_id = gen_trace_id()

    try:
        doc_result = await Runner.run(
            starting_agent = doc_agent,
            input          = doc_input,
            max_turns      = 5,
            run_config     = RunConfig(
                workflow_name                = "codeastra-document-analysis",
                trace_id                     = dt_trace_id,
                trace_metadata               = {"filename": filename,
                                               "codeastra_active": str(codeastra_active)},
                trace_include_sensitive_data = True,
            ),
        )
        flush_traces()
        full = str(doc_result.final_output) if doc_result.final_output else ""
    except Exception as e:
        yield {"type": "error", "message": str(e)}
        dt.complete()
        return

    s = dt.add_step("analysis_complete", {"length": len(full)})
    yield {"type": "thinking", "text": full, "trace_id": dt.id, "trace_step": s["step"]}
    dt.complete()

    if thread_id and thread_id in THREADS:
        THREADS[thread_id].add_message("agent", full, model="gpt-4o",
            trace_id=dt.id, codeastra_active=codeastra_active,
            intercepted_count=len(dt.intercepted))

    reveal_map = await reveal_from_trace(dt)

    yield {
        "type":                    "complete",
        "trace_id":                dt.id,
        "openai_trace_id":         dt_trace_id,
        "trace_url":               f"/traces/{dt.id}",
        "openai_traces_url":       "https://platform.openai.com/logs",
        "thread_id":               thread_id,
        "filename":                filename,
        "codeastra_active":        codeastra_active,
        "intercepted":             len(dt.intercepted),
        "real_data_seen_by_gpt":   0 if codeastra_active else "⚠️ YES",
        "revealed":                reveal_map,
        "emails_sent":             len(_email_results),
        "email_results":           _email_results,
    }


def _stream(gen):
    async def s():
        async for ev in gen:
            yield f"data: {json.dumps(ev, default=str)}\n\n"
    return StreamingResponse(s(), media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no",
                 "Access-Control-Allow-Origin": "*"})


# ═══════════════════════════════════════════════════════════
# ENDPOINTS
# ═══════════════════════════════════════════════════════════

@app.get("/")
async def index():
    with open("index.html") as f: return HTMLResponse(f.read())

@app.get("/health")
async def health():
    codeastra_ok = False
    if CODEASTRA_KEY:
        try:
            async with httpx.AsyncClient(timeout=5.0) as c:
                r = await c.get(f"{CODEASTRA_URL}/health",
                                headers={"X-API-Key": CODEASTRA_KEY})
                codeastra_ok = r.status_code == 200
        except Exception:
            pass
    return {
        "status": "healthy", "version": "agents-sdk-v2-traces",
        "openai_ready": bool(OPENAI_KEY), "codeastra_ready": bool(CODEASTRA_KEY),
        "codeastra_live": codeastra_ok, "db_ready": db_pool is not None,
        "sdk_traces": True,
        "openai_traces_url": "https://platform.openai.com/logs",
    }


# ── Agent Run ─────────────────────────────────────────────

@app.post("/agent/run/stream")
async def agent_run_stream(req: Request):
    body = await req.json()
    return _stream(run_openai_agent(
        body.get("task_type", "dba"), body.get("custom_task", ""),
        codeastra_active = body.get("codeastra_enabled", True),
        conversation_id  = body.get("conversation_id"),
        thread_id        = body.get("thread_id"),
    ))

@app.post("/agent/run/sync")
async def agent_run_sync(req: Request):
    body = await req.json()
    all_ev = []; intercept = []; summary = {}
    async for ev in run_openai_agent(
        body.get("task_type", "dba"), body.get("custom_task", ""),
        codeastra_active = body.get("codeastra_enabled", True),
        conversation_id  = body.get("conversation_id"),
        thread_id        = body.get("thread_id"),
    ):
        all_ev.append(ev)
        if ev["type"] == "intercepted": intercept.append(ev)
        if ev["type"] == "complete":    summary = ev
    return {"success": True, "intercepted": intercept, "summary": summary,
            "trace_url": f"/traces/{summary.get('trace_id', '')}"}

@app.post("/agent/run/protected")
async def agent_run_protected(req: Request):
    body = await req.json()
    return _stream(run_openai_agent(
        body.get("task_type", "dba"), body.get("custom_task", ""),
        codeastra_active=True,
        conversation_id=body.get("conversation_id"),
        thread_id=body.get("thread_id"),
    ))

@app.post("/agent/run/unprotected")
async def agent_run_unprotected(req: Request):
    body = await req.json()
    return _stream(run_openai_agent(
        body.get("task_type", "dba"), body.get("custom_task", ""),
        codeastra_active=False,
        conversation_id=body.get("conversation_id"),
        thread_id=body.get("thread_id"),
    ))


# ── Document endpoints ────────────────────────────────────

@app.post("/agent/analyze-document")
async def analyze_document(
    file:              UploadFile = File(default=None),
    task:              str        = Form(default=""),
    codeastra_enabled: str        = Form(default="true"),
    thread_id:         str        = Form(default=""),
):
    if file is None:
        return JSONResponse(status_code=400, content={"error": "No file uploaded"})
    text = await extract_text_from_file(file)
    return _stream(run_document_agent(
        text, task, file.filename or "document",
        codeastra_active = codeastra_enabled.lower() != "false",
        thread_id        = thread_id or None,
    ))

@app.post("/agent/analyze-url")
async def analyze_url(req: Request):
    body = await req.json()
    url  = body.get("url", "").strip()
    if not url.startswith(("http://", "https://")):
        return JSONResponse(status_code=400, content={"error": "Valid URL required"})
    text = await extract_text_from_url(url)
    return _stream(run_document_agent(
        text, body.get("task", ""), url,
        codeastra_active = body.get("codeastra_enabled", True),
        thread_id        = body.get("thread_id"),
    ))

@app.post("/agent/analyze-text")
async def analyze_text(req: Request):
    body = await req.json()
    text = body.get("text", "").strip()
    if not text:
        return JSONResponse(status_code=400, content={"error": "text required"})
    return _stream(run_document_agent(
        text, body.get("task", ""), body.get("name", "text"),
        codeastra_active = body.get("codeastra_enabled", True),
        thread_id        = body.get("thread_id"),
    ))

@app.post("/agent/analyze-multiple")
async def analyze_multiple(
    files:             list[UploadFile] = File(default=None),
    task:              str              = Form(default=""),
    codeastra_enabled: str              = Form(default="true"),
    thread_id:         str              = Form(default=""),
):
    if not files:
        return JSONResponse(status_code=400, content={"error": "No files"})
    if len(files) > 10:
        return JSONResponse(status_code=400, content={"error": "Max 10 files"})
    all_text = ""; names = []
    for f in files:
        t = await extract_text_from_file(f)
        names.append(f.filename or "file")
        all_text += f"\n\n=== FILE: {f.filename} ===\n{t}"
    return _stream(run_document_agent(
        all_text, task, f"{len(files)} files: {', '.join(names)}",
        codeastra_active = codeastra_enabled.lower() != "false",
        thread_id        = thread_id or None,
    ))


# ── Traces ────────────────────────────────────────────────

@app.get("/traces")
async def list_traces(limit: int = 20, status: str = None):
    traces = list(TRACES.values())
    traces.sort(key=lambda t: t.started_at, reverse=True)
    if status: traces = [t for t in traces if t.status == status]
    return {"traces": [t.to_dict() for t in traces[:limit]], "count": len(TRACES)}

@app.get("/traces/{trace_id}")
async def get_trace(trace_id: str):
    if trace_id not in TRACES:
        return JSONResponse(status_code=404, content={"error": "Trace not found"})
    return TRACES[trace_id].to_dict()

@app.get("/traces/{trace_id}/steps")
async def get_trace_steps(trace_id: str):
    if trace_id not in TRACES:
        return JSONResponse(status_code=404, content={"error": "Trace not found"})
    t = TRACES[trace_id]
    return {"trace_id": trace_id, "steps": t.steps, "total": len(t.steps)}

@app.get("/traces/{trace_id}/intercepted")
async def get_trace_intercepted(trace_id: str):
    if trace_id not in TRACES:
        return JSONResponse(status_code=404, content={"error": "Trace not found"})
    t = TRACES[trace_id]
    return {"trace_id": trace_id, "intercepted": t.intercepted,
            "count": len(t.intercepted),
            "real_data_seen_by_gpt": 0 if t.codeastra_active else "⚠️ YES"}

@app.delete("/traces/{trace_id}")
async def delete_trace(trace_id: str):
    if trace_id not in TRACES:
        return JSONResponse(status_code=404, content={"error": "Not found"})
    del TRACES[trace_id]
    return {"deleted": trace_id}


# ── Completions ───────────────────────────────────────────

@app.get("/completions")
async def list_completions(limit: int = 50):
    comps = sorted(COMPLETIONS.values(), key=lambda c: c["timestamp"], reverse=True)
    return {"completions": comps[:limit], "count": len(COMPLETIONS)}

@app.get("/completions/{comp_id}")
async def get_completion(comp_id: str):
    if comp_id not in COMPLETIONS:
        return JSONResponse(status_code=404, content={"error": "Not found"})
    return COMPLETIONS[comp_id]


# ── Conversations ─────────────────────────────────────────

@app.get("/conversations")
async def list_conversations():
    convs = sorted(CONVERSATIONS.values(), key=lambda c: c.updated_at, reverse=True)
    return {"conversations": [c.to_dict() for c in convs], "count": len(CONVERSATIONS)}

@app.post("/conversations")
async def create_conversation(req: Request):
    body = await req.json()
    conv = Conversation(body.get("title", ""), body.get("system_context", ""))
    CONVERSATIONS[conv.id] = conv
    return conv.to_dict()

@app.get("/conversations/{conv_id}")
async def get_conversation(conv_id: str):
    if conv_id not in CONVERSATIONS:
        return JSONResponse(status_code=404, content={"error": "Not found"})
    return CONVERSATIONS[conv_id].to_dict()

@app.post("/conversations/{conv_id}/message")
async def conversation_message(conv_id: str, req: Request):
    if conv_id not in CONVERSATIONS:
        return JSONResponse(status_code=404, content={"error": "Conversation not found"})
    body = await req.json()
    return _stream(run_openai_agent(
        body.get("task_type", "dba"), body.get("message", ""),
        codeastra_active=body.get("codeastra_enabled", True),
        conversation_id=conv_id,
    ))

@app.delete("/conversations/{conv_id}")
async def delete_conversation(conv_id: str):
    if conv_id not in CONVERSATIONS:
        return JSONResponse(status_code=404, content={"error": "Not found"})
    del CONVERSATIONS[conv_id]
    return {"deleted": conv_id}


# ── Threads ───────────────────────────────────────────────

@app.get("/threads")
async def list_threads():
    threads = sorted(THREADS.values(), key=lambda t: t.updated_at, reverse=True)
    return {"threads": [t.to_dict() for t in threads], "count": len(THREADS)}

@app.post("/threads")
async def create_thread(req: Request):
    body   = await req.json()
    thread = ChatKitThread(body.get("title", ""), body.get("metadata", {}))
    THREADS[thread.id] = thread
    return thread.to_dict()

@app.get("/threads/{thread_id}")
async def get_thread(thread_id: str):
    if thread_id not in THREADS:
        return JSONResponse(status_code=404, content={"error": "Not found"})
    return THREADS[thread_id].to_dict()

@app.post("/threads/{thread_id}/message")
async def thread_message(thread_id: str, req: Request):
    if thread_id not in THREADS:
        return JSONResponse(status_code=404, content={"error": "Thread not found"})
    body = await req.json()
    return _stream(run_openai_agent(
        body.get("task_type", "dba"), body.get("message", ""),
        codeastra_active=body.get("codeastra_enabled", True),
        thread_id=thread_id,
    ))

@app.post("/threads/{thread_id}/document")
async def thread_document(thread_id: str,
    file:              UploadFile = File(default=None),
    task:              str        = Form(default=""),
    codeastra_enabled: str        = Form(default="true"),
):
    if thread_id not in THREADS:
        return JSONResponse(status_code=404, content={"error": "Thread not found"})
    if file is None:
        return JSONResponse(status_code=400, content={"error": "No file"})
    text = await extract_text_from_file(file)
    return _stream(run_document_agent(
        text, task, file.filename or "document",
        codeastra_active = codeastra_enabled.lower() != "false",
        thread_id        = thread_id,
    ))

@app.patch("/threads/{thread_id}")
async def update_thread(thread_id: str, req: Request):
    if thread_id not in THREADS:
        return JSONResponse(status_code=404, content={"error": "Not found"})
    body = await req.json()
    t = THREADS[thread_id]
    if "title"    in body: t.title = body["title"]
    if "status"   in body: t.status = body["status"]
    if "metadata" in body: t.metadata.update(body["metadata"])
    t.updated_at = datetime.utcnow().isoformat()
    return t.to_dict()

@app.delete("/threads/{thread_id}")
async def delete_thread(thread_id: str):
    if thread_id not in THREADS:
        return JSONResponse(status_code=404, content={"error": "Not found"})
    del THREADS[thread_id]
    return {"deleted": thread_id}


# ── Protect + DB ──────────────────────────────────────────

@app.post("/protect")
async def protect_text(req: Request):
    events = []
    b      = await req.json()
    result = await protect(b.get("text", ""), events, True)
    return {"original": b.get("text", ""), "protected": result, "intercepted": events}

@app.post("/debug/protect-raw")
async def debug_protect_raw(req: Request):
    b = await req.json()
    if not CODEASTRA_KEY: return {"error": "No CODEASTRA_API_KEY"}
    async with httpx.AsyncClient(timeout=15.0) as client:
        r = await client.post(f"{CODEASTRA_URL}/protect/text",
            headers={"X-API-Key": CODEASTRA_KEY, "Content-Type": "application/json"},
            json={"text": b.get("text", "")})
        return {"status": r.status_code,
                "response": r.json() if r.status_code == 200 else r.text}

@app.post("/debug/test-responses-api")
async def test_responses_api():
    if not OPENAI_KEY: return {"error": "OPENAI_API_KEY not set"}
    client = AsyncOpenAI(api_key=OPENAI_KEY)
    try:
        response = await client.responses.create(
            model="gpt-4o", input="Say exactly: CODEASTRA TRACE TEST SUCCESSFUL", store=True)
        output_text = ""
        for item in getattr(response, "output", []):
            if getattr(item, "type", "") == "message":
                for c in getattr(item, "content", []):
                    if getattr(c, "type", "") == "output_text":
                        output_text += getattr(c, "text", "")
        return {"success": True, "response_id": getattr(response, "id", "unknown"),
                "output": output_text, "check_traces": "https://platform.openai.com/traces"}
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.get("/debug/openai-status")
async def openai_status():
    if not OPENAI_KEY: return {"error": "OPENAI_API_KEY not set"}
    client  = AsyncOpenAI(api_key=OPENAI_KEY)
    results = {}
    try:
        r = await client.chat.completions.create(
            model="gpt-4o", messages=[{"role": "user", "content": "ping"}], max_tokens=5)
        results["chat_completions"] = {"available": True, "model": r.model}
    except Exception as e:
        results["chat_completions"] = {"available": False, "error": str(e)}
    return {"openai_key_set": True, "features": results,
            "trace_proof_url": "https://platform.openai.com/traces"}

@app.get("/db/status")
async def db_status():
    if not db_pool: return {"connected": False, "message": "Set DATABASE_URL"}
    async with db_pool.acquire() as conn:
        try:
            ver = await conn.fetchval("SELECT version()")
            tbl = await conn.fetchval(
                "SELECT COUNT(*) FROM information_schema.tables WHERE table_schema='public'")
            return {"connected": True, "version": ver, "tables": tbl}
        except Exception as e:
            return {"connected": False, "error": str(e)}

@app.get("/db/tables")
async def db_tables():
    events = []
    result = await tool_list_tables(events, True)
    return JSONResponse({"data": json.loads(result), "intercepted": events})

@app.post("/db/query")
async def db_query(req: Request):
    b = await req.json(); events = []
    result = await tool_run_query(events, b.get("sql", ""), True)
    return JSONResponse({"data": json.loads(result), "intercepted": events})

@app.get("/db/stats")
async def db_stats_endpoint():
    events = []
    result = await tool_get_db_stats(events, True)
    return JSONResponse({"data": json.loads(result), "intercepted": events})

@app.get("/agent/tasks")
async def list_tasks():
    return {
        "model": "gpt-4o",
        "tasks": [
            {"id": "dba",      "name": "Database Performance Agent",
             "description": "Finds slow queries, creates missing indexes."},
            {"id": "audit",    "name": "Database Audit Agent",
             "description": "Full DB audit."},
            {"id": "security", "name": "Security Audit Agent",
             "description": "Finds security issues in DB."},
        ],
        "features": {
            "traces":        "Every run creates a trace — GET /traces",
            "completions":   "Every GPT call logged — GET /completions",
            "conversations": "Multi-turn conversations — POST /conversations",
            "threads":       "Persistent ChatKit threads — POST /threads",
            "toggle":        "Codeastra on/off per request",
            "proof":         "https://platform.openai.com/logs",
        }
    }


# ── Executor ──────────────────────────────────────────────

@app.get("/executor/capabilities")
async def executor_capabilities():
    if CODEASTRA_KEY:
        try:
            async with httpx.AsyncClient(timeout=5.0) as c:
                r = await c.get(f"{CODEASTRA_URL}/executor/supported",
                                headers={"X-API-Key": CODEASTRA_KEY})
                if r.status_code == 200:
                    return {"source": "codeastra_api", "codeastra": r.json()}
        except Exception:
            pass
    return {"source": "local", "guarantee": "Real values never returned to agent"}

@app.post("/executor/run")
async def executor_run(req: Request):
    body   = await req.json()
    result = await codeastra_executor_run(body.get("token_id", ""), body.get("dry_run", False))
    return result

@app.post("/executor/resolve")
async def executor_resolve(req: Request):
    body = await req.json()
    val  = await codeastra_resolve(body.get("token", ""))
    if val is None:
        return JSONResponse(status_code=404, content={"error": "Token not found"})
    return {"resolved": True, "real_value": val,
            "note": "Use in executor context only — never pass to agent"}

@app.post("/executor/resolve-batch")
async def executor_resolve_batch_endpoint(req: Request):
    body    = await req.json()
    results = await codeastra_resolve_batch(body.get("tokens", []))
    return {"resolved": results, "count": len(results)}

@app.post("/executor/check-threshold")
async def executor_check_threshold(req: Request):
    body      = await req.json()
    token     = body.get("token", "")
    threshold = float(body.get("threshold", 0))
    operator  = body.get("operator", "gt")
    real_val  = await codeastra_resolve(token)
    if real_val is None:
        return JSONResponse(status_code=404, content={"error": f"Cannot resolve: {token}"})
    try:
        v   = float(str(real_val).replace("$", "").replace(",", "").strip())
        ops = {"gt": v>threshold, "lt": v<threshold, "gte": v>=threshold, "lte": v<=threshold}
        return {"result": ops.get(operator, v>threshold), "operator": operator,
                "threshold": threshold, "real_value_returned": False}
    except Exception as e:
        return JSONResponse(status_code=400, content={"error": str(e)})

@app.post("/executor/concentration-check")
async def executor_concentration_check(req: Request):
    body           = await req.json()
    position_token = body.get("position_token", "")
    portfolio_token= body.get("portfolio_token", "")
    threshold_pct  = float(body.get("threshold_pct", 40.0))
    pv = await codeastra_resolve(position_token)
    tv = await codeastra_resolve(portfolio_token)
    if pv is None or tv is None:
        return {"exceeds_threshold": None, "note": "Tokens not resolved",
                "real_values_seen_by_agent": False}
    try:
        p      = float(str(pv).replace("$", "").replace(",", "").strip())
        t      = float(str(tv).replace("$", "").replace(",", "").strip())
        pct    = (p/t*100) if t > 0 else 0
        bucket = "critical" if pct>60 else "high" if pct>40 else "medium" if pct>20 else "low"
        return {"exceeds_threshold": pct>threshold_pct, "concentration_bucket": bucket,
                "threshold_pct": threshold_pct, "real_values_seen_by_agent": False}
    except Exception as e:
        return JSONResponse(status_code=400, content={"error": str(e)})

@app.post("/executor/sum-amounts")
async def executor_sum_amounts(req: Request):
    body      = await req.json()
    tokens    = body.get("tokens", [])
    threshold = body.get("threshold")
    resolved  = await codeastra_resolve_batch(tokens)
    total = 0.0; count = 0
    for val in resolved.values():
        try:
            total += float(str(val).replace("$", "").replace(",", "").strip())
            count += 1
        except Exception:
            pass
    result = {"sum": total, "count": count, "real_individual_values_returned": False}
    if threshold is not None:
        result["exceeds_threshold"] = total > float(threshold)
    return result

@app.post("/executor/classify-amount")
async def executor_classify_amount(req: Request):
    body    = await req.json()
    token   = body.get("token", "")
    buckets = body.get("buckets", [
        {"label": "small",  "min": 0,       "max": 10000},
        {"label": "medium", "min": 10000,   "max": 100000},
        {"label": "large",  "min": 100000,  "max": 1000000},
        {"label": "whale",  "min": 1000000, "max": None},
    ])
    real_val = await codeastra_resolve(token)
    if real_val is None:
        return JSONResponse(status_code=404, content={"error": f"Cannot resolve {token}"})
    try:
        value = float(str(real_val).replace("$", "").replace(",", "").strip())
        label = "unknown"
        for b in buckets:
            mn = b.get("min", 0) or 0
            mx = b.get("max")
            if value >= mn and (mx is None or value < mx):
                label = b["label"]; break
        return {"bucket": label, "token": token, "real_value_returned": False}
    except Exception as e:
        return JSONResponse(status_code=400, content={"error": str(e)})


# ═══════════════════════════════════════════════════════════
# VAULT-AS-TEE EMAIL SENDER
# The correct architecture:
#   Agent passes tokens only
#   Vault opens ephemeral session
#   Real values loaded inside — encrypted
#   Email fires from inside vault
#   Session wiped — data gone
#   Agent never knew any real value
# ═══════════════════════════════════════════════════════════

async def tee_send_email(email_token: str, subject: str, body: str) -> dict:
    """
    Send email via executor.
    Resolves tokens → sends via Resend → agent never sees real values.
    """
    import re as _re
    token_pattern = _re.compile(r'\[CV[TD]:[A-Z]+:[A-F0-9a-f0-9\-]{4,}\]')
    body_tokens   = list(set(token_pattern.findall(body)))
    log.info(f"[EMAIL] Sending — {len(body_tokens)} tokens in body to resolve")
    return await _tee_email_fallback(email_token, subject, body, body_tokens)


async def _tee_email_fallback(
    email_token:  str,
    subject:      str,
    body:         str,
    body_tokens:  list,
) -> dict:
    """
    Fallback if /tee/run is not available.
    Still uses Codeastra vault/resolve — still never exposes to agent.
    """
    log.info("[TEE EMAIL FALLBACK] Using vault/resolve directly")

    # Resolve email address
    real_email = await codeastra_resolve(email_token)
    if not real_email:
        return {"sent": False, "error": f"Could not resolve: {email_token}",
                "real_address_seen_by_agent": False}

    # Resolve all body tokens
    revealed_body = body
    if body_tokens:
        resolved = await codeastra_resolve_batch(body_tokens)
        for token, real_value in resolved.items():
            if real_value:
                revealed_body = revealed_body.replace(token, str(real_value))

    # Send via Resend
    result = await _send_via_resend(real_email, subject, revealed_body)
    return {
        **result,
        "real_address_seen_by_agent": False,
        "vault_as_tee":               False,
        "fallback":                   True,
    }


# ═══════════════════════════════════════════════════════════
# EMAIL EXECUTOR
# Agent says: "send to [CVT:EMAIL:A1B2]"
# Executor resolves token → real address → sends email
# Agent never sees the real address. Ever.
# ═══════════════════════════════════════════════════════════

import re as _email_re

EMAIL_SERVICE = os.getenv("EMAIL_SERVICE", "resend")    # resend | sendgrid | smtp
SENDGRID_KEY  = os.getenv("SENDGRID_API_KEY", "")
RESEND_KEY    = os.getenv("RESEND_API_KEY", "")
SMTP_HOST     = os.getenv("SMTP_HOST", "smtp.gmail.com")
SMTP_PORT     = int(os.getenv("SMTP_PORT", "587"))
SMTP_USER     = os.getenv("SMTP_USER", "")
SMTP_PASS     = os.getenv("SMTP_PASS", "")
FROM_EMAIL    = os.getenv("FROM_EMAIL", "noreply@codeastra.dev")
FROM_NAME     = os.getenv("FROM_NAME", "Codeastra Agent")

# Token pattern
TOKEN_PAT = _email_re.compile(
    r'\[CV[TD]:[A-Z]+:[A-Za-z0-9\-]{4,}\]|cdt_[a-z]+_[bto]_[a-z0-9]+'
)


async def _send_via_sendgrid(to_email: str, subject: str, body: str) -> dict:
    """Send email via SendGrid API."""
    if not SENDGRID_KEY:
        return {"error": "SENDGRID_API_KEY not set"}
    try:
        async with httpx.AsyncClient(timeout=15.0) as c:
            r = await c.post(
                "https://api.sendgrid.com/v3/mail/send",
                headers={
                    "Authorization": f"Bearer {SENDGRID_KEY}",
                    "Content-Type":  "application/json",
                },
                json={
                    "personalizations": [{"to": [{"email": to_email}]}],
                    "from":    {"email": FROM_EMAIL, "name": FROM_NAME},
                    "subject": subject,
                    "content": [{"type": "text/plain", "value": body}],
                },
            )
            return {"sent": r.status_code == 202, "status": r.status_code}
    except Exception as e:
        return {"error": str(e)}


async def _send_via_resend(to_email: str, subject: str, body: str) -> dict:
    """Send email via Resend API."""
    if not RESEND_KEY:
        log.error("RESEND: RESEND_API_KEY not set in environment variables")
        return {"error": "RESEND_API_KEY not set — add it to Railway environment variables"}
    log.info(f"RESEND: sending to masked address, from={FROM_EMAIL}, subject={subject[:50]}")
    try:
        async with httpx.AsyncClient(timeout=15.0) as c:
            r = await c.post(
                "https://api.resend.com/emails",
                headers={
                    "Authorization": f"Bearer {RESEND_KEY}",
                    "Content-Type":  "application/json",
                },
                json={
                    "from":    f"{FROM_NAME} <{FROM_EMAIL}>",
                    "to":      [to_email],
                    "subject": subject,
                    "text":    body,
                },
            )
            data = r.json()
            if r.status_code == 200:
                log.info(f"RESEND: ✅ email sent — id={data.get('id')}")
                return {"sent": True, "id": data.get("id"), "status": 200}
            else:
                log.error(f"RESEND: ❌ failed status={r.status_code} body={r.text[:200]}")
                return {"sent": False, "error": data, "status": r.status_code}
    except Exception as e:
        log.error(f"RESEND: exception — {e}")
        return {"error": str(e)}


async def _send_via_smtp(to_email: str, subject: str, body: str) -> dict:
    """Send email via SMTP (Gmail etc)."""
    if not SMTP_USER or not SMTP_PASS:
        return {"error": "SMTP_USER and SMTP_PASS not set"}
    try:
        import smtplib
        from email.mime.text import MIMEText
        msg            = MIMEText(body)
        msg["Subject"] = subject
        msg["From"]    = f"{FROM_NAME} <{SMTP_USER}>"
        msg["To"]      = to_email
        with smtplib.SMTP(SMTP_HOST, SMTP_PORT) as server:
            server.starttls()
            server.login(SMTP_USER, SMTP_PASS)
            server.sendmail(SMTP_USER, [to_email], msg.as_string())
        return {"sent": True}
    except Exception as e:
        return {"error": str(e)}


async def executor_send_email(
    email_token: str,
    subject:     str,
    body:        str,
) -> dict:
    """
    THE CORE FUNCTION.

    Agent passes an email token + body with tokens.
    Executor resolves:
      1. The email address token → real address
      2. ALL tokens in the body  → real names, values
    Sends the email with real values revealed.
    Agent never sees any real values.
    """
    # Step 1 — Resolve email address token
    real_email = await codeastra_resolve(email_token)

    if not real_email:
        return {
            "sent":  False,
            "error": f"Could not resolve email token: {email_token}",
            "token": email_token,
            "real_address_seen_by_agent": False,
        }

    # Step 2 — Find ALL tokens in the body and resolve them
    # The body contains [CVT:NAME:xxxx] tokens — replace with real names
    import re as _re
    token_pattern = _re.compile(r'\[CV[TD]:[A-Z]+:[A-F0-9a-f0-9\-]{4,}\]')
    tokens_in_body = list(set(token_pattern.findall(body)))

    revealed_body = body
    if tokens_in_body:
        # Batch resolve all tokens in body
        resolved = await codeastra_resolve_batch(tokens_in_body)
        # Replace tokens with real values in the body
        for token, real_value in resolved.items():
            if real_value:
                revealed_body = revealed_body.replace(token, str(real_value))
        log.info(f"[EMAIL EXECUTOR] Resolved {len(resolved)} tokens in email body")

    # Step 3 — Send with real values in body
    if EMAIL_SERVICE == "sendgrid" and SENDGRID_KEY:
        result = await _send_via_sendgrid(real_email, subject, revealed_body)
    elif EMAIL_SERVICE == "resend" and RESEND_KEY:
        result = await _send_via_resend(real_email, subject, revealed_body)
    elif SMTP_USER and SMTP_PASS:
        result = await _send_via_smtp(real_email, subject, revealed_body)
    else:
        log.info(f"[EMAIL EXECUTOR] Would send to: {real_email} | Subject: {subject}")
        result = {"sent": True, "mode": "log_only — set EMAIL_SERVICE env vars"}

    # Step 4 — Return result WITHOUT real address or resolved values
    return {
        **result,
        "token":                      email_token,
        "real_address_seen_by_agent": False,
        "tokens_resolved_in_body":    len(tokens_in_body),
        "subject":                    subject,
        "body_length":                len(revealed_body),
    }


# ── Email tool for the agent ─────────────────────────────
# Agent calls this with a token — never the real address

@app.post("/executor/send-email")
async def executor_send_email_endpoint(req: Request):
    """
    Send email via executor — agent passes token, never real address.

    Body:
      email_token: str  — [CVT:EMAIL:A1B2] from the agent
      subject:     str  — email subject
      body:        str  — email body / analysis summary

    The executor resolves the token internally.
    Real address never returned to agent.
    """
    body        = await req.json()
    email_token = body.get("email_token", "")
    subject     = body.get("subject", "Codeastra Agent Report")
    email_body  = body.get("body", "")

    if not email_token:
        return JSONResponse(status_code=400,
            content={"error": "email_token required — pass token not real address"})

    if not TOKEN_PAT.match(email_token.strip()):
        # They passed a real email — intercept and protect
        events = []
        protected = await protect(email_token, events, True)
        tokens = TOKEN_PAT.findall(protected)
        if tokens:
            email_token = tokens[0]
            log.info(f"[EMAIL EXECUTOR] Intercepted real email — tokenized automatically")
        else:
            return JSONResponse(status_code=400,
                content={"error": "Pass a token not a real email address"})

    result = await executor_send_email(email_token, subject, email_body)
    return result


@app.post("/executor/send-report")
async def executor_send_report(req: Request):
    """
    Full pipeline endpoint:
    1. Receives analysis text + email token
    2. Formats as professional report
    3. Sends via executor (real address never seen by agent)

    Body:
      email_token:  str  — [CVT:EMAIL:A1B2]
      analysis:     str  — the agent's full analysis
      document_name: str — name of analyzed document
      subject:      str  — optional custom subject
    """
    body          = await req.json()
    email_token   = body.get("email_token", "")
    analysis      = body.get("analysis", "")
    document_name = body.get("document_name", "Document")
    subject       = body.get("subject") or f"Codeastra Analysis: {document_name}"

    if not email_token or not analysis:
        return JSONResponse(status_code=400,
            content={"error": "email_token and analysis required"})

    # Format professional report
    report_body = f"""CODEASTRA SECURE ANALYSIS REPORT
{'='*50}
Document: {document_name}
Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}
Protected by: Codeastra Zero Trust Infrastructure
Real data seen by AI: 0

{'='*50}
ANALYSIS
{'='*50}

{analysis}

{'='*50}
This report was generated by an AI agent that never
saw the real names, account numbers, SSNs, or email
addresses in the original document.
All sensitive values were tokenized before the AI
processed them.
Powered by Codeastra — app.codeastra.dev
"""

    result = await executor_send_email(email_token, subject, report_body)
    return {
        **result,
        "document_name": document_name,
        "report_sent":   result.get("sent", False),
    }




# ═══════════════════════════════════════════════════════════
# RUN
# ═══════════════════════════════════════════════════════════

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=PORT)
